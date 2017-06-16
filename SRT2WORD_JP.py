import re, io, sys, os
from docx import Document
cwd = os.getcwd()

for file in os.listdir(cwd):
    if file.endswith('.srt'):
        files = file
        if sys.argv[1] in files:
            f = io.open(sys.argv[1], encoding='utf-8', mode='r+')
            regex = re.compile('(?<=\d{2}:\d{2}:\d{2},\d{3}\s-->\s\d{2}:\d{2}:\d{2},\d{3}\n).*')
            subtitles = regex.findall(f.read())

            #write subtitles into a TXT file
            txt = open('demo.txt', 'a')
            for subtitle in subtitles:
                subtitle = subtitle.replace('\\N', '')
                regexN = re.compile('す$', re.MULTILINE)
                subtitle = re.sub(regexN, 'す\n', subtitle)
                txt.write(subtitle)
            txt.close()


document = Document()
document.add_heading(sys.argv[1] + '\n', 3)
f = open('demo.txt', 'r')
lines = f.readlines()
for line in lines:
    document.add_paragraph(line)

document.save(sys.argv[1] + '.docx')

os.remove('demo.txt')


